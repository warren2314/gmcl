from __future__ import annotations

import os
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "ineligible-guide" / "GMCL_Ineligible_Player_Staff_Guide.docx"
TMP = ROOT / "tmp" / "ineligible-guide"

RED = "C41E3A"
RED_DARK = "8E1428"
BLUE = "2E86C1"
GOLD = "F5A623"
INK = "17212B"
MUTED = "5F6B76"
CANVAS = "F4F6F8"
BORDER = "DFE4E8"
GREEN = "198754"
GREEN_LIGHT = "EAF7F0"
AMBER = "FFF4D6"
BLUE_LIGHT = "EAF5FB"
RED_LIGHT = "FCECEF"
WHITE = "FFFFFF"
BLACK = "000000"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def set_para(p, before=0, after=6, line=1.15, keep_next=False, keep_together=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together


def clear_cell(cell):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        p._element.remove(run._element)
    return p


def add_cell_text(cell, text, *, bold=False, color=INK, size=10, align=None):
    p = clear_cell(cell)
    if align is not None:
        p.alignment = align
    set_para(p, after=0, line=1.1)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def add_field(run, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_alt_text(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, *, bold_lead=None, color=INK, after=6, italic=False):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.2)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead):])
        set_run(r2, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, color=color, italic=italic)
    return p


def add_list(doc, items, ordered=False, level=0):
    style = "GMCL Number" if ordered else "GMCL Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Inches(0.38 + level * 0.24)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        set_para(p, after=4, line=1.2)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            set_run(r1, bold=True)
            r2 = p.add_run(rest)
            set_run(r2)
        else:
            r = p.add_run(item)
            set_run(r)


def add_callout(doc, label, text, *, kind="info"):
    palette = {
        "info": (BLUE_LIGHT, BLUE),
        "success": (GREEN_LIGHT, GREEN),
        "warning": (AMBER, GOLD),
        "danger": (RED_LIGHT, RED),
        "plain": (CANVAS, MUTED),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": accent}, top={"val": "nil"}, right={"val": "nil"}, bottom={"val": "nil"})
    p = clear_cell(cell)
    set_para(p, after=0, line=1.15)
    r = p.add_run(label.upper() + "  ")
    set_run(r, size=9, bold=True, color=accent)
    r = p.add_run(text)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_status_chip(cell, text, fill, color=WHITE):
    set_cell_shading(cell, fill)
    p = add_cell_text(cell, text, bold=True, color=color, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_data_table(doc, headers, rows, widths, *, header_fill=RED, first_col_bold=False, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], header_fill)
        add_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE if header_fill != CANVAS else INK, size=9)
    border = {"val": "single", "sz": "4", "color": BORDER}
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            add_cell_text(cells[idx], str(value), bold=(first_col_bold and idx == 0), size=font_size)
            set_cell_border(cells[idx], top=border, bottom=border, left=border, right=border)
    set_table_geometry(table, widths)
    return table


def add_screen_panel(doc, title, subtitle, rows):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, WHITE)
    border = {"val": "single", "sz": "8", "color": BORDER}
    set_cell_border(cell, top=border, bottom=border, left=border, right=border)
    p = clear_cell(cell)
    set_para(p, after=2)
    r = p.add_run(title)
    set_run(r, size=11, bold=True, color=RED)
    p = cell.add_paragraph()
    set_para(p, after=6)
    r = p.add_run(subtitle)
    set_run(r, size=9, color=MUTED)
    for label, value in rows:
        p = cell.add_paragraph()
        set_para(p, after=3, line=1.1)
        r = p.add_run(label + ": ")
        set_run(r, size=9, bold=True, color=INK)
        r = p.add_run(value)
        set_run(r, size=9, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_email_preview(doc, label, audience, subject, body_lines, *, warning=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, WHITE)
    border = {"val": "single", "sz": "8", "color": BORDER}
    set_cell_border(cell, top=border, bottom=border, left=border, right=border)
    p = clear_cell(cell)
    set_para(p, after=2)
    r = p.add_run("EMAIL PREVIEW - " + label.upper())
    set_run(r, size=9, bold=True, color=RED)
    p = cell.add_paragraph()
    set_para(p, after=2)
    r = p.add_run("To: ")
    set_run(r, size=9, bold=True)
    r = p.add_run(audience)
    set_run(r, size=9)
    p = cell.add_paragraph()
    set_para(p, after=8)
    r = p.add_run("Subject: ")
    set_run(r, size=9, bold=True)
    r = p.add_run(subject)
    set_run(r, size=9)
    for line in body_lines:
        p = cell.add_paragraph()
        set_para(p, after=5, line=1.15)
        r = p.add_run(line)
        set_run(r, size=9.5)
    if warning:
        p = cell.add_paragraph()
        set_para(p, before=5, after=0)
        r = p.add_run("CHECK BEFORE APPROVAL: ")
        set_run(r, size=8.5, bold=True, color=GOLD)
        r = p.add_run(warning)
        set_run(r, size=8.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_step_header(doc, number, title, purpose, time_text):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1200, 8160])
    left, right = table.rows[0].cells
    set_cell_shading(left, RED)
    set_cell_shading(right, CANVAS)
    add_cell_text(left, f"STEP {number}", bold=True, color=WHITE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = clear_cell(right)
    set_para(p, after=1)
    r = p.add_run(title)
    set_run(r, size=17, bold=True, color=INK)
    p = right.add_paragraph()
    set_para(p, after=0)
    r = p.add_run(purpose + "  |  Typical time: " + time_text)
    set_run(r, size=9.5, color=MUTED)
    return table


def page_break(doc):
    doc.add_page_break()


def make_image_png(source: Path, target: Path):
    target.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(source) as image:
        image.convert("RGBA").save(target, "PNG")


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element

    def add_num(abstract_id, num_id, fmt, text, left=540, hanging=270, color=INK):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "288")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        clr = OxmlElement("w:color")
        clr.set(qn("w:val"), color)
        r_pr.append(clr)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_num(80, 80, "bullet", "-", color=RED)
    add_num(81, 81, "decimal", "%1.", color=RED)
    styles = doc.styles
    for style_name, num_id in (("GMCL Bullet", 80), ("GMCL Number", 81)):
        if style_name not in styles:
            style = styles.add_style(style_name, 1)
        else:
            style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        p_pr = style.element.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numid = OxmlElement("w:numId")
        numid.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, numid])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    specs = {
        "Title": (30, RED, 0, 8),
        "Subtitle": (14, MUTED, 0, 14),
        "Heading 1": (18, RED, 18, 8),
        "Heading 2": (14, BLUE, 14, 7),
        "Heading 3": (11.5, INK, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(doc, logo_png: Path):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.34)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [1800, 7560], indent=0)
    table.cell(0, 0).width = Inches(1.25)
    p = clear_cell(table.cell(0, 0))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    shape = run.add_picture(str(logo_png), width=Inches(0.92))
    add_alt_text(shape, "GMCL logo", "Greater Manchester Cricket League logo")
    p = clear_cell(table.cell(0, 1))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("INELIGIBLE PLAYER CASEWORK | STAFF GUIDE")
    set_run(r, size=8, bold=True, color=MUTED)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("GMCL staff guide | August 2026   |   Page ")
    set_run(r, size=8, color=MUTED)
    field_run = p.add_run()
    set_run(field_run, size=8, color=MUTED)
    add_field(field_run, "PAGE")


def add_cover(doc, logo_png, hawk_png):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=36, after=18)
    shape = p.add_run().add_picture(str(logo_png), width=Inches(2.35))
    add_alt_text(shape, "GMCL logo", "Greater Manchester Cricket League logo")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=8)
    r = p.add_run("INELIGIBLE PLAYER CASEWORK")
    set_run(r, size=25, bold=True, color=RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=18)
    r = p.add_run("A simple staff walkthrough and usability proof")
    set_run(r, size=15, bold=True, color=BLUE)

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2500, 6860])
    left, right = table.rows[0].cells
    set_cell_shading(left, CANVAS)
    p = clear_cell(left)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(hawk_png), width=Inches(1.25))
    add_alt_text(shape, "Hawk AI mascot", "GMCL Hawk AI cricket hawk mascot")
    set_cell_shading(right, CANVAS)
    p = clear_cell(right)
    set_para(p, after=5)
    r = p.add_run("Designed for the person doing the work")
    set_run(r, size=13, bold=True, color=INK)
    p = right.add_paragraph()
    set_para(p, after=4)
    r = p.add_run("Short instructions, clear next actions, plain-English help and safe email previews.")
    set_run(r, size=10.5, color=INK)
    p = right.add_paragraph()
    set_para(p, after=0)
    r = p.add_run("Staff guide v1.0 | Fictional case data | Training copy")
    set_run(r, size=9, color=MUTED, italic=True)

    add_callout(doc, "Important", "This guide explains the improved staff experience. Human approval, privacy checks and a permanent case history remain in place. All names, clubs, dates and addresses in the example are fictional.", kind="warning")
    add_body(doc, "Example used throughout: GMCL-2026-0042, a fictional report concerning Jordan Taylor, Riverside CC 2nd XI v Northbridge CC 2nd XI on 2 August 2026.", color=MUTED, italic=True, after=0)


def build_document():
    TMP.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    logo_png = TMP / "gmcl-logo.png"
    hawk_png = TMP / "hawk-ai.png"
    make_image_png(ROOT / "images" / "logo.webp", logo_png)
    make_image_png(ROOT / "images" / "hawk-ai-mascot.webp", hawk_png)

    doc = Document()
    configure_styles(doc)
    configure_numbering(doc)
    configure_page(doc, logo_png)

    # Cover
    add_cover(doc, logo_png, hawk_png)
    page_break(doc)

    # Quick start
    add_heading(doc, "Start here: the two-minute version", 1)
    add_callout(doc, "Purpose", "Receive a report, check it, raise or link a case, investigate fairly, approve the outcome independently, then send only the correct information to each audience.", kind="info")
    add_heading(doc, "The safe route through the work", 2)
    flow = doc.add_table(rows=2, cols=6)
    set_table_geometry(flow, [1560] * 6)
    labels = ["1. Review", "2. Map", "3. Raise case", "4. Investigate", "5. Approve", "6. Send"]
    details = ["Check the report", "Confirm clubs, team, date and player", "Create or link", "Evidence, rule and response", "Different staff member", "Preview then publish"]
    for idx in range(6):
        set_cell_shading(flow.cell(0, idx), RED if idx in (0, 2, 5) else BLUE)
        add_cell_text(flow.cell(0, idx), labels[idx], bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(flow.cell(1, idx), CANVAS)
        add_cell_text(flow.cell(1, idx), details[idx], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Three rules to remember", 2)
    add_list(doc, [
        ("Nothing is decided automatically. ", "A report or Hawk AI suggestion starts staff review; it does not prove a breach."),
        ("Preview before sending. ", "Every email and attachment must be visible before it is approved or queued."),
        ("Green means the case exists. ", "Once an intake is linked to a case, its queue card becomes green and shows the case reference."),
    ])
    add_heading(doc, "What the improved feature provides", 2)
    add_data_table(doc, ["Area", "What staff now see", "Why it helps"], [
        ("Intake and cases", "Clear next actions and a green Case raised state", "You can see immediately what is finished and what to do next"),
        ("Queue", "My Work / All Work tabs and newest/oldest date order", "Focus on your own work or cover the full queue"),
        ("Excel import", "Verified runs and rows hidden by default", "Completed checks leave the working list but remain in history"),
        ("Email", "One preview centre for every saved message and audience", "Review the exact recipients, subject and body before sending"),
        ("HawkAI", "Cited eligibility-rule suggestions prefilled for review", "Staff get a useful starting point while keeping the final decision"),
    ], [1700, 3700, 3960], font_size=8.7)
    add_callout(doc, "Plain English", "Technical labels are replaced where possible. When a technical control matters, the screen explains it in everyday language; the glossary at the end gives extra detail.", kind="success")
    page_break(doc)

    # Work queue
    add_heading(doc, "Your home screen", 1)
    add_body(doc, "The home screen answers one question immediately: What do I need to do next?")
    add_screen_panel(doc, "INELIGIBLE PLAYER WORK", "Current queue layout", [
        ("Tabs", "My Work (default) | All Work"),
        ("Sort", "Newest first (default) | Oldest first"),
        ("Filters", "Needs review | Case raised | Waiting for club | Decision | Complete | Exceptions"),
        ("Help", "A next-action sentence appears on every card"),
    ])
    add_heading(doc, "What each colour means", 2)
    status = doc.add_table(rows=5, cols=3)
    set_table_geometry(status, [1500, 2600, 5260])
    for i, text in enumerate(("Colour", "Meaning", "What you do")):
        set_cell_shading(status.cell(0, i), RED)
        add_cell_text(status.cell(0, i), text, bold=True, color=WHITE, size=9)
    rows = [
        ("RED", RED, "Needs action", "Open it and follow the next-action helper."),
        ("AMBER", GOLD, "Waiting or needs checking", "Check the outstanding item or due date."),
        ("GREEN", GREEN, "Case raised or task complete", "Open the linked case if further work is required."),
        ("GREY", MUTED, "History / no current action", "No action unless you deliberately reopen history."),
    ]
    for idx, (label, fill, meaning, action) in enumerate(rows, 1):
        add_status_chip(status.cell(idx, 0), label, fill)
        add_cell_text(status.cell(idx, 1), meaning, bold=True, size=9)
        add_cell_text(status.cell(idx, 2), action, size=9)
    add_callout(doc, "Example", "The fictional Jordan Taylor report appears under My Work, sorted near the top because it was received recently. Its card says: Next: confirm the offending team and fixture date.", kind="info")
    add_heading(doc, "My Work and All Work", 2)
    add_list(doc, [
        ("My Work: ", "shows cases assigned to you."),
        ("All Work: ", "shows the whole queue for cover, supervision and workload balancing."),
        ("Keep the date order: ", "moving between My Work and All Work keeps your chosen newest/oldest direction."),
    ])
    page_break(doc)

    # Step 1
    add_step_header(doc, 1, "Review a new report", "Decide whether the report has enough information to progress", "3-5 minutes")
    add_heading(doc, "What you should see", 2)
    add_screen_panel(doc, "INTAKE IPR-000142", "Received 7 August 2026 at 09:18 | Assigned to you", [
        ("Reporting club", "Northbridge CC"),
        ("Offending club/team", "Riverside CC / Riverside 2nd XI"),
        ("Player", "Jordan Taylor"),
        ("Fixture", "2 August 2026"),
        ("Reason", "Reported appearance may conflict with the starred-player restriction"),
        ("Next action", "Confirm the mapped team, player and fixture before raising a case"),
    ])
    add_heading(doc, "Do this", 2)
    add_list(doc, [
        "Read the report and open each evidence item.",
        "Confirm the selected club and team are the correct database records.",
        "Confirm the player name and fixture date match the evidence.",
        "Choose Create new case, Link to existing case, Duplicate or Ignore.",
    ], ordered=True)
    add_callout(doc, "Helper text", "Not sure which option to choose? Create new case starts a separate investigation. Link to existing case adds this report to work already under way. Duplicate records that the same report was received again. Ignore requires a reason and keeps an audit record.", kind="info")
    add_heading(doc, "Done when", 2)
    add_body(doc, "You have selected a route and the screen shows a confirmation. If you raised or linked a case, the intake card is green and displays its case reference.", bold_lead="You have selected a route")
    page_break(doc)

    # Step 2 Excel
    add_step_header(doc, 2, "Review the Excel import", "Reconcile historical tracker rows without resending or recalculating old outcomes", "varies by batch")
    add_callout(doc, "Safety", "Uploading, reviewing and verifying the spreadsheet does not send an email or create a new sanction.", kind="success")
    add_heading(doc, "Simple views", 2)
    add_data_table(doc, ["View", "What it contains", "Default?"], [
        ("Needs checking", "Rows that still need a match, decision or explanation, including exceptions", "Yes"),
        ("Verified history", "Rows already checked", "No"),
    ], [2100, 5400, 1860], first_col_bold=True)
    add_heading(doc, "Do this", 2)
    add_list(doc, [
        "Upload the approved tracker file and check the filename, sheet name and row count shown on screen.",
        "Work through Needs verification. Confirm the matched intake and whether the historical case was open or closed.",
        "Explain any free-text points or cards entry before marking the row verified.",
        "Ask the named signatory to review the summary and apply the completed batch.",
    ], ordered=True)
    add_callout(doc, "After verification", "The row disappears from the working list immediately. It remains available under Show verified history, with reviewer, date, source row and checksum details for audit purposes.", kind="success")
    add_screen_panel(doc, "IMPORT SUMMARY", "Example after two rows are verified", [
        ("Needs verification", "8"),
        ("Exceptions", "1"),
        ("Verified history", "32 (hidden)"),
        ("Next action", "Resolve the one exception before requesting sign-off"),
    ])
    page_break(doc)

    # Step 3
    add_step_header(doc, 3, "Raise or link the case", "Move a checked intake into the investigation workflow", "2-4 minutes")
    add_heading(doc, "Create a new case", 2)
    add_list(doc, [
        "Select the offending team from the mapped list.",
        "Check the fixture date and player name.",
        "Use the public allegation field for wording that may later be shared with the offending club.",
        "Keep reporter contact details and private investigation notes in the private field only.",
        "Select Raise case.",
    ], ordered=True)
    add_callout(doc, "Green confirmation", "Once the case is created, the intake becomes green and reads: Case raised - GMCL-2026-0042. A prominent Open case button replaces the Create case button.", kind="success")
    add_screen_panel(doc, "CASE RAISED", "Green status card", [
        ("Reference", "GMCL-2026-0042"),
        ("Status", "Investigating"),
        ("Assigned to", "A. Investigator"),
        ("Next action", "Review Hawk AI's rule suggestion and retained scorecard evidence"),
    ])
    add_heading(doc, "Link to an existing case", 2)
    add_body(doc, "Search by case reference, club, team or player. The system shows the current case status before you confirm. The intake still turns green because the case link now exists.")
    add_callout(doc, "What does linked mean?", "This report is now attached to the named case. The original report is not changed or deleted.", kind="plain")
    page_break(doc)

    # Step 4 Hawk
    add_step_header(doc, 4, "Review the Hawk AI rule suggestion", "Turn evidence into a cited draft rule allegation for staff review", "2-5 minutes")
    add_callout(doc, "How HawkAI works here", "HawkAI searches GMCL's own active published-rule index for eligibility wording and pre-fills likely rule candidates inside the case. It does not send the player, club or case summary outside GMCL.", kind="warning")
    add_screen_panel(doc, "HAWK AI SUGGESTION", "Draft only - staff confirmation required", [
        ("Suggested rule", "Rule 3.5 - Starred players"),
        ("Why it may apply", "The active rule contains eligibility or ineligible-player wording"),
        ("Source", "Published GMCL rules, Rule 3.5 [Open source]"),
        ("Staff check", "Compare the cited wording with the player, team and fixture evidence"),
        ("Actions", "Refresh suggestions | Edit | Save reviewed rule"),
    ])
    add_heading(doc, "Do this", 2)
    add_list(doc, [
        "Open each rule source cited by Hawk AI.",
        "Compare the rule wording with the case evidence and fixture date.",
        "Use or edit the suggested allegation. Never accept it only because HawkAI placed it first.",
        "Save the reviewed alleged rule before requesting a club response.",
    ], ordered=True)
    add_callout(doc, "Human decision", "Hawk AI suggests; staff decide. It must never create a finding, sanction, email or case automatically.", kind="danger")
    add_heading(doc, "Done when", 2)
    add_body(doc, "The case displays Reviewed rule allegation with the staff member, date, rule source and exact saved wording.")
    page_break(doc)

    # Step 5 evidence
    add_step_header(doc, 5, "Check the evidence", "Make sure the investigation is based on the correct fixture and source material", "5-10 minutes")
    add_heading(doc, "Evidence checklist", 2)
    add_list(doc, [
        "Open the retained report and any uploaded files.",
        "Confirm the Play-Cricket match ID, date and team.",
        "Use Fetch latest scorecard if no snapshot is present or a later version is required.",
        "Check that the named player appears on the correct side of the scorecard.",
        "Record a private investigation note explaining what the evidence establishes and what remains uncertain.",
    ], ordered=True)
    add_callout(doc, "Why the system may stop", "If more than one fixture matches, the system will not guess. Correct the mapped team or match ID, then try again.", kind="warning")
    add_screen_panel(doc, "PLAY-CRICKET SCORECARD EVIDENCE", "Private evidence", [
        ("Match ID", "7458963"),
        ("Fixture", "Riverside CC 2nd XI v Northbridge CC 2nd XI"),
        ("Date", "2 August 2026"),
        ("Players retained", "22"),
        ("Integrity", "Stored copy checked against its digital fingerprint"),
    ])
    add_callout(doc, "Digital fingerprint", "A short code created from a file's contents. If the file changes, the code changes too. Technical name: checksum or SHA-256.", kind="plain")
    page_break(doc)

    # Step 6 email response
    add_step_header(doc, 6, "Preview and request the club response", "Show the exact proposed email before anything is queued", "5 minutes")
    add_heading(doc, "Email preview centre", 2)
    add_body(doc, "The Email preview centre shows every latest saved message for the case, including its audience, status, recipients, subject and full wording.")
    add_data_table(doc, ["Message", "Audience", "Status", "Action"], [
        ("Response request", "Offending club", "Ready to preview", "Review and queue"),
        ("Day-five reminder", "Offending club", "Prepared", "Preview only"),
        ("Outcome", "Offending club", "Not available", "Created after approval"),
        ("Outcome", "Reporting club", "Not available", "Created after approval"),
        ("Official outcome", "League recipients", "Not available", "Created after approval"),
    ], [2500, 2100, 1900, 2860], font_size=8.5)
    add_email_preview(doc, "Response request", "secretary@riverside.example.invalid", "Response requested for GMCL case GMCL-2026-0042", [
        "Dear Club Secretary,",
        "The GMCL requests an official response from Riverside CC 2nd XI concerning the following allegation:",
        "Jordan Taylor may have appeared in the fixture on 2 August 2026 contrary to the reviewed starred-player restriction.",
        "Alleged rule: Rule 3.5 - Starred players. This is an allegation for response, not a finding.",
        "Use the secure link below to respond and upload supporting evidence. [Secure response link]",
        "Case reference: GMCL-2026-0042. The link expires seven days after this email is delivered.",
        "Regards, Greater Manchester Cricket League",
    ], warning="No reporter name, email, telephone number or reporting-club identity appears in the message.")
    page_break(doc)

    add_heading(doc, "Response reminder preview", 1)
    add_email_preview(doc, "Day-five reminder", "secretary@riverside.example.invalid", "Reminder: Response requested for GMCL case GMCL-2026-0042", [
        "Dear Club Secretary,",
        "This is the single reminder that the response for GMCL case GMCL-2026-0042 is due in two days.",
        "No adverse decision is made automatically if the deadline passes.",
        "[Secure response link]",
        "Regards, Greater Manchester Cricket League",
    ])
    add_heading(doc, "Before you queue the request", 2)
    add_list(doc, [
        "The recipient is the verified official mailbox for the offending club.",
        "The public allegation is accurate and neutral.",
        "The reviewed alleged rule is included.",
        "Reporter and reporting-club identity are absent.",
        "Any evidence being shared is a separately redacted copy approved for sharing.",
    ])
    add_callout(doc, "Clock starts on delivery", "The seven-day response window and day-five reminder begin only after the first email is accepted by the mail provider.", kind="info")
    page_break(doc)

    # Step 7 response
    add_step_header(doc, 7, "Review the club response", "Add the response to the case and record what it changes", "5-15 minutes")
    add_heading(doc, "When a response arrives", 2)
    add_list(doc, [
        "Open the New response task in My Work.",
        "Read the response and any uploaded evidence.",
        "Mark each item reviewed and record a private investigation note.",
        "If the response arrived by email, phone or meeting, use Add manual response so it appears on the same timeline.",
        "Continue investigating or move to Propose decision when the evidence is complete.",
    ], ordered=True)
    add_callout(doc, "No response", "At day seven the response link closes and the case returns to investigation. The system does not decide that the allegation is true.", kind="warning")
    add_screen_panel(doc, "CASE TIMELINE", "All entries are dated and attributed", [
        ("7 Aug", "Allegation and rule reviewed by A. Investigator"),
        ("8 Aug", "Response request delivered; deadline 15 Aug"),
        ("13 Aug", "Reminder delivered"),
        ("14 Aug", "Club response received with one attachment"),
        ("14 Aug", "Response reviewed; investigation continues"),
    ])
    page_break(doc)

    # Step 8 decision
    add_step_header(doc, 8, "Propose and approve the decision", "Record findings, rule and effects with independent approval", "10-20 minutes")
    add_heading(doc, "The proposer records", 2)
    add_list(doc, [
        "Confirmed findings in audience-safe wording.",
        "The reviewed rule determination, or an explicit not-applicable determination.",
        "One complete bundle of effects: for example warning, points adjustment, fine, ban, card or no action.",
        "Effective dates and appeal instructions where required.",
        "Private reasoning separately from wording that may be sent externally.",
    ])
    add_screen_panel(doc, "DUMMY DECISION", "For illustration only - not a real finding", [
        ("Finding", "The appearance was confirmed after review of the scorecard and club response"),
        ("Rule", "Rule 3.5 - Starred players"),
        ("Effect", "Warning (dummy example)"),
        ("Proposed by", "A. Investigator"),
        ("Required next action", "Independent approval by another authorised staff member"),
    ])
    add_heading(doc, "The approver checks", 2)
    add_list(doc, [
        "The proposer and approver are different people.",
        "The latest intake revision has been reviewed.",
        "The findings match the evidence and response.",
        "Each effect applies to the correct player, team or match.",
        "All recipient groups and email/PDF previews are correct and privacy-safe.",
    ])
    add_callout(doc, "Technical term", "Independent approval is sometimes called separation of duties. It means the person who proposes the outcome cannot approve their own work.", kind="plain")
    page_break(doc)

    # Step 9 outcomes
    add_step_header(doc, 9, "Preview and publish the outcome", "Send the approved wording to the correct audiences", "5-10 minutes")
    add_callout(doc, "Preview all", "Use the Email preview centre to review each saved audience message. Use the outcome-preview buttons to open the offending-club, reporting-club and official PDFs before issuing an approved outcome.", kind="info")
    add_email_preview(doc, "Offending-club outcome", "secretary@riverside.example.invalid", "GMCL ineligible-player case outcome GMCL-2026-0042", [
        "Dear Club Secretary,",
        "The independently approved decision for case GMCL-2026-0042 is set out below.",
        "Findings: [approved findings]",
        "Rule determination: [approved rule wording]",
        "Decision and sanctions: [approved effect summary]",
        "Appeal instructions: [approved instructions]",
        "Regards, Greater Manchester Cricket League",
    ], warning="This audience may receive the full approved findings and sanctions, but never reporter identity or private notes.")
    page_break(doc)

    add_heading(doc, "Other outcome previews", 1)
    add_email_preview(doc, "Reporting-club outcome", "secretary@northbridge.example.invalid", "GMCL ineligible-player case outcome GMCL-2026-0042", [
        "Dear Club Secretary,",
        "The GMCL has completed its investigation into the ineligible-player report recorded as case GMCL-2026-0042.",
        "Confirmed findings: [approved findings]",
        "Rule determination: [approved rule wording]",
        "Final outcome: [approved outcome or no-action result]",
        "Regards, Greater Manchester Cricket League",
    ], warning="Do not include the offending club's private response, internal notes, private evidence or reporter details.")
    add_email_preview(doc, "Official outcome", "Configured Executive and discipline recipients", "GMCL ineligible-player case outcome GMCL-2026-0042", [
        "Approved league outcome record",
        "Case: GMCL-2026-0042 | Source: ineligible_player",
        "Offending club: Riverside CC | Reporting club: Northbridge CC",
        "Findings, rule determination, decision, sanctions and appeal instructions follow in the approved record.",
    ], warning="Play-Cricket recipients are added for league-table points; finance recipients are added for fines.")
    add_callout(doc, "No action", "A no-action decision uses the same independent approval and audience previews. It closes without appearing on the public register.", kind="success")
    page_break(doc)

    # Completion and corrections
    add_heading(doc, "After publication", 1)
    add_data_table(doc, ["What happens", "What staff see", "Owner"], [
        ("Messages queued", "Email status for every recipient", "System / email operations"),
        ("Points adjustment", "Two-day Play-Cricket task", "Configured points administrator"),
        ("Fine", "Finance follow-up", "Finance recipient"),
        ("No action", "Case closed, unpublished", "Case team"),
        ("All work complete", "Green Complete status", "Assigned investigator"),
    ], [3000, 3800, 2560], font_size=9)
    add_heading(doc, "If new information arrives after approval", 2)
    add_body(doc, "The system may offer Reopen for source change only if no outcome may have been sent, nothing has been published and no follow-up task has started.")
    add_list(doc, [
        "Read the newer intake revision and enter a clear reopening reason.",
        "The old approved record stays visible; the system adds a correction record instead of deleting history.",
        "Merge the new revision, investigate again, propose a new decision and obtain independent approval again.",
    ], ordered=True)
    add_callout(doc, "Why history is kept", "A reliable case record shows what was known and decided at each point in time. Keeping history protects clubs, staff and the league.", kind="plain")
    page_break(doc)

    # Glossary
    add_heading(doc, "Plain-English glossary", 1)
    add_body(doc, "Preferred screen wording appears first. The technical term is included only so staff can recognise it in audit or support discussions.")
    glossary = [
        ("Case history cannot be changed", "Immutable", "The original record is kept. Corrections are added as new entries rather than overwriting it."),
        ("New report", "Intake", "A report waiting to be checked. It is not yet a finding or sanction."),
        ("Saved version", "Revision", "A new copy created when source information changes."),
        ("Digital fingerprint", "Checksum / SHA-256", "A code used to confirm that a stored file has not changed."),
        ("Initial review", "Triage", "Checking what the report is, whether it is complete and where it should go."),
        ("Source history", "Provenance", "Where information came from, when it was collected and which version was used."),
        ("Personal details removed", "Redacted", "A safe copy with private information removed before sharing."),
        ("Waiting to send", "Outbox", "An approved email held for the delivery service."),
        ("Different person approves", "Separation of duties", "The proposer cannot approve their own decision."),
        ("Stop safely when uncertain", "Fail closed", "The system pauses and asks a person to resolve the problem instead of guessing."),
    ]
    add_data_table(doc, ["Use on screen", "Technical term", "Meaning"], glossary, [2600, 2100, 4660], first_col_bold=True, font_size=8.8)
    add_callout(doc, "Writing rule", "Use the everyday phrase in buttons and headings. Put the technical term in a tooltip, Learn more panel or audit detail - not in the main instruction.", kind="info")
    page_break(doc)

    # Troubleshooting
    add_heading(doc, "When the system will not let you continue", 1)
    trouble = [
        ("I cannot raise a case", "Confirm the team mapping, fixture date, player and configured assignee."),
        ("The scorecard cannot be found", "Check the Play-Cricket team ID and fixture date. The system will not choose between multiple matches."),
        ("I cannot request a response", "Record the alleged rule, save both email drafts, verify the official mailbox and confirm outbound email is enabled."),
        ("I cannot propose a decision", "Close or expire the response window and merge the newest linked intake revision."),
        ("I cannot approve", "The approver must be authorised and must not be the proposer."),
        ("I cannot publish", "Check privacy warnings, recipient mailboxes and approved email/PDF snapshots."),
        ("A verified Excel row is missing", "Open Show verified history; verified rows are hidden from the working list by design."),
        ("Hawk AI suggests the wrong rule", "Choose Edit or Mark not applicable, record why, and report feedback from the source panel."),
    ]
    add_data_table(doc, ["Problem", "What to check"], trouble, [3000, 6360], first_col_bold=True, font_size=9)
    add_heading(doc, "Ask for help with useful details", 2)
    add_list(doc, [
        "Case or intake reference.",
        "The action you were trying to complete.",
        "The exact message shown on screen.",
        "Do not copy reporter contact details into general support channels.",
    ])
    page_break(doc)

    # Acceptance map
    add_heading(doc, "Usability acceptance checklist", 1)
    add_body(doc, "This section turns the eight feedback points into observable outcomes for design review and user acceptance testing.")
    acceptance = [
        ("1", "Simpler workflow", "Every card has one next action; each case shows a step tracker; help is available without leaving the page."),
        ("2", "Plain English", "Technical terms are removed from primary UI or explained inline and in a glossary."),
        ("3", "Case raised is green", "Creating or linking a case immediately turns the intake card green and shows Open case."),
        ("4", "Date order", "Queue supports Newest first and Oldest first and remembers the choice."),
        ("5", "Verified Excel rows", "Verified rows disappear from the default working list and remain under verified history."),
        ("6", "See every email", "A preview centre shows recipient, subject, body, attachments, status and all outcome audiences before sending."),
        ("7", "My work / all work", "My Work is the default tab; All Work is available to permitted staff."),
        ("8", "Hawk AI assists", "Hawk AI pre-fills a cited draft rule allegation, asks for missing facts and requires staff confirmation."),
    ]
    add_data_table(doc, ["#", "Requirement", "Pass condition"], acceptance, [600, 2500, 6260], font_size=8.7)
    add_heading(doc, "Suggested user test", 2)
    add_body(doc, "Give a staff tester the fictional Jordan Taylor intake without coaching. Ask them to raise the case, review HawkAI's rule suggestion, preview the response request and explain what happens next. The feature succeeds if they can complete the route and explain the safety checks in their own words.")
    add_callout(doc, "Final check", "A new user should never need to understand database tables, hashes, immutable records, rollout gates or outbox processing to complete normal casework.", kind="success")

    # Core properties
    doc.core_properties.title = "GMCL Ineligible Player Casework - Staff Guide"
    doc.core_properties.subject = "End-user walkthrough, training example and usability acceptance guide"
    doc.core_properties.author = "Greater Manchester Cricket League"
    doc.core_properties.keywords = "GMCL, ineligible player, staff guide, Hawk AI, casework"
    doc.core_properties.comments = "Fictional example data. GMCL staff training guide."

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
