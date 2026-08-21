from __future__ import annotations

import os
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = Path(
    os.environ.get(
        "GMCL_INELIGIBLE_QUICK_GUIDE_OUT",
        ROOT / "outputs" / "ineligible-guide" / "GMCL_Ineligible_Player_Quick_Guide.docx",
    )
)
TMP = ROOT / "tmp" / "ineligible-quick-guide"

# compact_reference_guide preset tokens.
PAGE_WIDTH = 8.5
PAGE_HEIGHT = 11.0
MARGIN = 1.0
HEADER_FOOTER_DISTANCE = 0.492
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120
BASE_FONT = "Calibri"
BODY_SIZE = 11
BODY_AFTER = 6
BODY_LINE = 1.25
H1_SIZE, H1_BEFORE, H1_AFTER = 16, 18, 10
H2_SIZE, H2_BEFORE, H2_AFTER = 13, 14, 7
H3_SIZE, H3_BEFORE, H3_AFTER = 12, 10, 5
LIST_LEFT_DXA = 540
LIST_HANGING_DXA = 270
LIST_AFTER_TWIPS = 80
LIST_LINE_TWIPS = 300

BLUE = "2E74B5"
NAVY = "1F4D78"
INK = "20252B"
MUTED = "66717D"
LIGHT_BLUE = "E8EEF5"
BORDER = "D7DBE2"
WHITE = "FFFFFF"
PALE_GREEN = "E9F5EC"  # Named success-state override.
GREEN = "397A45"
PALE_AMBER = "FFF4D6"  # Named warning-state override.
AMBER = "A66B00"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, size=BODY_SIZE, bold=False, color=INK, italic=False):
    run.font.name = BASE_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BASE_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BASE_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def set_paragraph(p, *, before=0, after=BODY_AFTER, line=BODY_LINE, keep_next=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep_next


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("end", CELL_START_END_DXA),
    ):
        node = tc_mar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, color=BORDER, size="6", sides=("top", "left", "bottom", "right")):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in sides:
        edge = borders.find(qn("w:" + side))
        if edge is None:
            edge = OxmlElement("w:" + side)
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), color)


def clear_cell(cell):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        p._element.remove(run._element)
    return p


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"table widths must total {CONTENT_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(BODY_AFTER)
    normal.paragraph_format.line_spacing = BODY_LINE

    for name, size, color, before, after in (
        ("Heading 1", H1_SIZE, BLUE, H1_BEFORE, H1_AFTER),
        ("Heading 2", H2_SIZE, BLUE, H2_BEFORE, H2_AFTER),
        ("Heading 3", H3_SIZE, NAVY, H3_BEFORE, H3_AFTER),
    ):
        style = doc.styles[name]
        style.font.name = BASE_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = BODY_LINE
        style.paragraph_format.keep_with_next = True


def add_numbering_abstract(doc, abstract_id: int, *, ordered: bool):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(LIST_LEFT_DXA))
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(LIST_LEFT_DXA))
    indent.set(qn("w:hanging"), str(LIST_HANGING_DXA))
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), str(LIST_AFTER_TWIPS))
    spacing.set(qn("w:line"), str(LIST_LINE_TWIPS))
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BASE_FONT)
    fonts.set(qn("w:hAnsi"), BASE_FONT)
    r_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)


def configure_numbering(doc):
    add_numbering_abstract(doc, 90, ordered=True)
    add_numbering_abstract(doc, 91, ordered=False)


_next_num_id = 200


def new_num_id(doc, abstract_id: int) -> int:
    global _next_num_id
    _next_num_id += 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(_next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    doc.part.numbering_part.element.append(num)
    return _next_num_id


def apply_num(p, num_id: int):
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])


def add_rich_text(p, segments):
    for text, bold in segments:
        run = p.add_run(text)
        set_run_font(run, bold=bold)


def add_list(doc, items, *, ordered=True):
    num_id = new_num_id(doc, 90 if ordered else 91)
    for item in items:
        p = doc.add_paragraph()
        set_paragraph(p, after=4, line=BODY_LINE)
        apply_num(p, num_id)
        if isinstance(item, str):
            add_rich_text(p, [(item, False)])
        else:
            add_rich_text(p, item)


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text, *, bold_lead=None, italic=False, color=INK, after=BODY_AFTER):
    p = doc.add_paragraph()
    set_paragraph(p, after=after)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic, color=color)
    return p


def add_callout(doc, label, text, *, kind="info"):
    fill, accent = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (PALE_GREEN, GREEN),
        "warning": (PALE_AMBER, AMBER),
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size="18", sides=("left",))
    p = clear_cell(cell)
    set_paragraph(p, after=0)
    run = p.add_run(label.upper() + "  ")
    set_run_font(run, size=9, bold=True, color=accent)
    run = p.add_run(text)
    set_run_font(run)
    spacer = doc.add_paragraph()
    set_paragraph(spacer, after=0, line=1.0)
    return table


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        p = clear_cell(cell)
        set_paragraph(p, after=0, line=1.0)
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_border(cell)
            p = clear_cell(cell)
            set_paragraph(p, after=0, line=1.15)
            run = p.add_run(value)
            set_run_font(run, size=10, bold=index == 1)
        set_table_geometry(table, widths)
    return table


def add_field(run, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def configure_page(doc, logo_png: Path):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH)
    section.page_height = Inches(PAGE_HEIGHT)
    section.top_margin = Inches(MARGIN)
    section.bottom_margin = Inches(MARGIN)
    section.left_margin = Inches(MARGIN)
    section.right_margin = Inches(MARGIN)
    section.header_distance = Inches(HEADER_FOOTER_DISTANCE)
    section.footer_distance = Inches(HEADER_FOOTER_DISTANCE)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [1250, 8110])
    left, right = table.rows[0].cells
    p = clear_cell(left)
    set_paragraph(p, after=0, line=1.0)
    p.add_run().add_picture(str(logo_png), width=Inches(0.62))
    p = clear_cell(right)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(p, after=0, line=1.0)
    run = p.add_run("GMCL  |  INELIGIBLE-PLAYER WORK")
    set_run_font(run, size=8, bold=True, color=MUTED)
    for cell in table.rows[0].cells:
        set_cell_border(cell, color=BORDER, size="6", sides=("bottom",))

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(p, after=0, line=1.0)
    run = p.add_run("GMCL quick guide  |  August 2026  |  Page ")
    set_run_font(run, size=8, color=MUTED)
    field = p.add_run()
    set_run_font(field, size=8, color=MUTED)
    add_field(field, "PAGE")


def add_title_stack(doc):
    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=4, line=1.0)
    run = p.add_run("GMCL STAFF QUICK GUIDE")
    set_run_font(run, size=9, bold=True, color=BLUE)
    p = doc.add_paragraph()
    set_paragraph(p, after=6, line=1.0, keep_next=True)
    run = p.add_run("Ineligible-player work")
    set_run_font(run, size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_paragraph(p, after=16, line=1.15)
    run = p.add_run("Three routes. One controlled case process.")
    set_run_font(run, size=14, color=MUTED)

    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [2700, 6660])
    for row_index, (label, value) in enumerate(
        (("Audience", "GMCL casework staff"), ("Purpose", "Raise, import and reconcile safely"))
    ):
        left, right = table.rows[row_index].cells
        for cell in (left, right):
            set_cell_border(cell)
        set_cell_shading(left, LIGHT_BLUE)
        p = clear_cell(left)
        set_paragraph(p, after=0)
        set_run_font(p.add_run(label), size=10, bold=True, color=NAVY)
        p = clear_cell(right)
        set_paragraph(p, after=0)
        set_run_font(p.add_run(value), size=10)
    spacer = doc.add_paragraph()
    set_paragraph(spacer, after=0, line=1.0)


def page_break(doc):
    doc.add_page_break()


def add_flow_arrow(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=2, line=1.0)
    run = p.add_run("\u2193")
    set_run_font(run, size=16, bold=True, color=BLUE)
    return p


def audit_document(doc):
    section = doc.sections[0]
    assert section.page_width == Inches(PAGE_WIDTH)
    assert section.page_height == Inches(PAGE_HEIGHT)
    assert section.top_margin == Inches(MARGIN)
    assert section.bottom_margin == Inches(MARGIN)
    assert section.left_margin == Inches(MARGIN)
    assert section.right_margin == Inches(MARGIN)
    assert section.header_distance.twips == round(HEADER_FOOTER_DISTANCE * 1440)
    assert section.footer_distance.twips == round(HEADER_FOOTER_DISTANCE * 1440)

    normal = doc.styles["Normal"]
    assert normal.font.name == BASE_FONT and normal.font.size == Pt(BODY_SIZE)
    assert normal.paragraph_format.space_after == Pt(BODY_AFTER)
    assert normal.paragraph_format.line_spacing == BODY_LINE
    for name, size, before, after in (
        ("Heading 1", H1_SIZE, H1_BEFORE, H1_AFTER),
        ("Heading 2", H2_SIZE, H2_BEFORE, H2_AFTER),
        ("Heading 3", H3_SIZE, H3_BEFORE, H3_AFTER),
    ):
        style = doc.styles[name]
        assert style.font.name == BASE_FONT and style.font.size == Pt(size)
        assert style.paragraph_format.space_before == Pt(before)
        assert style.paragraph_format.space_after == Pt(after)
        assert style.paragraph_format.line_spacing == BODY_LINE

    tables = list(doc.tables) + list(section.header.tables)
    for table in tables:
        tbl_pr = table._tbl.tblPr
        assert tbl_pr.find(qn("w:tblW")).get(qn("w:w")) == str(CONTENT_DXA)
        assert tbl_pr.find(qn("w:tblInd")).get(qn("w:w")) == str(TABLE_INDENT_DXA)
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid_widths) == CONTENT_DXA
        for row in table.rows:
            assert len(row.cells) == len(grid_widths)
            for index, cell in enumerate(row.cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                assert tc_pr.find(qn("w:tcW")).get(qn("w:w")) == str(grid_widths[index])
                tc_mar = tc_pr.find(qn("w:tcMar"))
                assert tc_mar.find(qn("w:top")).get(qn("w:w")) == str(CELL_TOP_BOTTOM_DXA)
                assert tc_mar.find(qn("w:start")).get(qn("w:w")) == str(CELL_START_END_DXA)

    numbered = 0
    for paragraph in doc.paragraphs:
        p_pr = paragraph._p.pPr
        if p_pr is not None and p_pr.find(qn("w:numPr")) is not None:
            numbered += 1
    assert numbered >= 30

    numbering = doc.part.numbering_part.element
    abstract_ids = {
        node.get(qn("w:abstractNumId")) for node in numbering.findall(qn("w:abstractNum"))
    }
    assert {"90", "91"}.issubset(abstract_ids)


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    logo_png = TMP / "gmcl-logo.png"
    with Image.open(ROOT / "images" / "logo.webp") as image:
        image.convert("RGBA").save(logo_png, "PNG")

    doc = Document()
    configure_styles(doc)
    configure_numbering(doc)
    configure_page(doc, logo_png)

    # Page 1: choose the route.
    add_title_stack(doc)
    add_callout(
        doc,
        "Important",
        "Importing is not the same as raising a case. The Google Form import brings reports in for staff to choose and review. The tracker applies historical information only. Staff deliberately raise every live case.",
        kind="warning",
    )
    add_heading(doc, "Start here", 1)
    add_list(
        doc,
        [
            [("Sign in to the GMCL admin portal.", False)],
            [("Click ", False), ("Sanctions", True), (".", False)],
            [("Click ", False), ("Ineligible-player work", True), (".", False)],
            [("On ", False), ("Ineligible-player cases", True), (", choose the route that matches your task.", False)],
        ],
    )
    add_table(
        doc,
        ["What do you need to do?", "Choose"],
        [
            ("Review one report and raise its case", "Route 1 - Raise one case"),
            ("Import reports and choose which to progress", "Route 2 - Import and choose reports"),
            ("Reconcile the approved historical workbook", "Route 3 - Import historical tracker"),
        ],
        [4750, 4610],
    )
    add_heading(doc, "What 'raise one case manually' means", 2)
    add_body(
        doc,
        "An ineligible-player case must start from a private report in the queue. Do not use Add card, ban, fine or points decision to create a blank ineligible-player case.",
    )
    add_callout(
        doc,
        "If the report is missing",
        "Complete the current private Google Form, run Route 2, then open the report from the queue.",
        kind="info",
    )
    page_break(doc)

    # Page 2: raise a live case.
    add_heading(doc, "Route 1 - Raise one case", 1)
    add_list(
        doc,
        [
            [("Click ", False), ("Sanctions", True), (", then ", False), ("Ineligible-player work", True), (".", False)],
            [("Under ", False), ("Route 1 - Raise one case", True), (", click ", False), ("Open next selected report", True), (".", False)],
            [("If the button says ", False), ("View reports", True), (", click it, then click ", False), ("Review report", True), (" beside the correct entry.", False)],
            [("Check ", False), ("Reported details", True), (" and any evidence.", False)],
            [("Under ", False), ("Raise this case", True), (", check Offending team, Reporting club, Fixture date and Player.", False)],
            [("If needed, click ", False), ("Review case wording", True), (" and check the recorded allegation and private investigation context.", False)],
            [("Click ", False), ("Raise case", True), (".", False)],
            [("When Case [reference] is ready appears, click ", False), ("Open case", True), (".", False)],
        ],
    )
    add_callout(doc, "Expected result", "An investigation opens. No email is sent and no outcome is decided.", kind="success")
    add_heading(doc, "If a new case should not be raised", 2)
    add_list(
        doc,
        [
            [("Existing investigation: ", True), ("click Other outcomes, complete Link to an existing case, then click Link and merge intake.", False)],
            [("Duplicate report: ", True), ("complete Resolve without a new case, then click Mark duplicate.", False)],
            [("Irrelevant or no action: ", True), ("record the reason, then click Ignore intake.", False)],
        ],
        ordered=False,
    )
    page_break(doc)

    # Page 3: import and choose the active work list.
    add_heading(doc, "Route 2 - Import and choose reports", 1)
    add_callout(
        doc,
        "Values, not colours",
        "Import the full configured Google response sheet. Do not delete, copy or reorder live response rows; choose the active work on the next screen.",
        kind="warning",
    )
    add_list(
        doc,
        [
            "Make sure the required reports are present in the private Google Form response sheet.",
            [("Click ", False), ("Sanctions", True), (", then ", False), ("Ineligible-player work", True), (".", False)],
            [("Under Route 2, click ", False), ("Import and choose reports", True), (".", False)],
            "Read the summary. Source rows read is every Google response; added and changed are database updates, so zero is normal on a repeat import.",
            [("If the page says ", False), ("Selection is blocked", True), (", record the import number and error, plus any spreadsheet row shown, then stop for import or identity help.", False)],
            [("Use ", False), ("Fixture from", True), (", Fixture to, Order or search to find the current handover.", False)],
            [("Tick ", False), ("Progress", True), (" or use Select all shown. Filtering never unticks a report already selected.", False)],
            [("Check the selected total against the current handover, enter a label such as ", False), ("Dave handover - 11 Aug 2026", True), (", then save.", False)],
            [("Click ", False), ("Review report", True), (" beside the first selected report and follow Route 1.", False)],
        ],
    )
    add_heading(doc, "If a report looks missing", 2)
    add_list(
        doc,
        [
            "Clear the search and fixture-date filters.",
            "Open Report history and search for the player.",
            "Do not select the report again if it is already case-raised, linked, duplicate or ignored.",
            "Treat one spreadsheet row as one report, even when its Player box contains several names.",
        ],
        ordered=True,
    )
    add_callout(doc, "Expected result", "Selected open reports appear in the normal queue. Unselected open reports are hidden, not deleted. Progressed reports remain in Report history.", kind="success")
    add_callout(doc, "Need attention", "A row warning does not automatically mean the report is missing. Read the warning shown beside that report; only an identity or manifest problem blocks selection.", kind="info")
    page_break(doc)
    # Page 4: tracker upload and row review.
    add_heading(doc, "Route 3 - Import the historical tracker", 1)
    add_callout(
        doc,
        "Use only the approved workbook",
        "The file must be .xlsx, no larger than 16 MB, with the Form responses 1 sheet and expected columns A to Z.",
        kind="warning",
    )
    add_heading(doc, "Step 1 - Upload", 2)
    add_list(
        doc,
        [
            [("Click ", False), ("Sanctions", True), (", then ", False), ("Ineligible-player work", True), (".", False)],
            [("Under Route 3, click ", False), ("Open tracker import", True), (".", False)],
            [("Under Step 1, click ", False), ("Tracker (.xlsx, max 16 MB)", True), (".", False)],
            "Select the approved workbook.",
            [("Click ", False), ("Upload tracker", True), (".", False)],
        ],
    )
    add_body(doc, "The system opens Import check #[number].", bold_lead="The system opens")
    add_heading(doc, "Step 2 - Check every row", 2)
    add_list(
        doc,
        [
            [("Leave ", False), ("Needs checking", True), (" selected.", False)],
            [("Review the totals for ", False), ("Needs checking", True), (", Suggested matches and Needs help.", False)],
            "Check the player, club, team, suggested Google intake and historical state on each row.",
            [("For a correct straightforward suggestion, click ", False), ("Confirm suggested match", True), (".", False)],
            [("Repeat until every straightforward row has moved to ", False), ("Verified history", True), (".", False)],
        ],
    )
    add_heading(doc, "When a row needs the full review", 2)
    add_body(doc, "If Confirm suggested match is not shown:")
    add_list(
        doc,
        [
            "Check Reconciliation and the Google intake ID.",
            "Choose the Historical case state.",
            "Complete the Points/cards review if it appears.",
            "Record the review reason and reviewer.",
            [("Click ", False), ("Verify row", True), (".", False)],
        ],
    )
    add_callout(doc, "Do not guess", "Ask the casework lead to review ambiguous matches or points/cards wording.", kind="warning")
    page_break(doc)

    # Page 5: tracker completion and downstream case process.
    add_heading(doc, "Finish the tracker import", 1)
    add_heading(doc, "Step 3 - Sign off", 2)
    add_list(
        doc,
        [
            [("When Needs checking reaches zero, find ", False), ("Step 3 - Sign off", True), (".", False)],
            "Check Your name and read the confirmation statement.",
            "Tick Save my name and confirmation in the audit history.",
            [("Click ", False), ("Sign off import", True), (".", False)],
        ],
    )
    add_heading(doc, "Step 4 - Apply history", 2)
    add_list(
        doc,
        [
            "Review the application summary and Application note.",
            "Tick the one-time application confirmation.",
            [("Click ", False), ("Apply signed-off history", True), (".", False)],
        ],
    )
    add_callout(
        doc,
        "Tracker result",
        "Signed-off private history and reviewed open/closed status are applied once. Unmatched or excluded rows remain untouched.",
        kind="success",
    )
    add_callout(
        doc,
        "What the tracker cannot do",
        "It cannot create a case, decision, sanction, points/cards entry, task, correspondence or email.",
        kind="warning",
    )
    add_heading(doc, "Send the first email after raising a case", 1)
    add_callout(doc, "Where to start", "Use the large Next action: contact the club for its explanation section. Work down the three numbered cards in order.", kind="info")
    add_list(
        doc,
        [
            "If prompted, record and save the alleged rule under investigation.",
            "Review the initial email, then select Save initial email. Saving does not contact the club.",
            "Optional: tick the safe-test confirmation and select Send TEST copy to me. It goes only to your administrator email, has no live response link and does not change the case.",
            "Review the reminder, then select Save reminder. It is prepared now but is not sent now.",
            "Check the displayed To address, then select Send initial email to club only when you intend to contact the club.",
            "Check delivery history. The seven-day period starts after delivery; the reminder is sent on day five only if still needed.",
        ],
    )
    add_heading(doc, "Assign the case or ask for help", 2)
    add_list(doc, [
        "Under Case owner and help, choose an administrator, enter the reason and select Save case owner to hand over the whole investigation.",
        "To keep ownership, choose another administrator under Give a supporting task, describe the work, add an optional due date and select Assign supporting task.",
        "Open supporting tasks remain visible on the case and in the task list. Every change is audited.",
    ], ordered=False)
    add_heading(doc, "Safety checks", 2)
    add_list(
        doc,
        [
            "Raising a case does not send an email.",
            "Saving either draft does not send an email. A TEST copy goes only to your administrator address and cannot open a club response window.",
            "Send initial email to club is the deliberate live-send action.",
            "A decision requires approval by a different authorised administrator.",
            "Outcomes are not sent until Denver selects Final sign-off and issue outcomes.",
        ],
        ordered=False,
    )

    page_break(doc)

    # Page 6: board-level controlled flow.
    add_heading(doc, "Board flow - controlled intake to outcome", 1)
    add_body(
        doc,
        "Every source row is accounted for, but only open and unlinked reports are offered for selection.",
    )
    add_callout(
        doc,
        "1  SAFE IMPORT",
        "Read every configured Google Form response into the private intake area. Nothing is deleted and no case or email is created.",
        kind="info",
    )
    add_flow_arrow(doc)
    add_table(
        doc,
        ["OPEN + UNLINKED", "ALREADY PROGRESSED", "IDENTITY NEEDS HELP"],
        [
            (
                "Available to choose",
                "Retained in report history or its case",
                "Selection blocks for manual checking",
            ),
        ],
        [3120, 3120, 3120],
    )
    add_flow_arrow(doc)
    add_callout(
        doc,
        "2  STAFF SELECTION",
        "Staff tick the exact open reports to progress and save the work list.",
        kind="info",
    )
    add_flow_arrow(doc)
    add_table(
        doc,
        ["SELECTED", "NOT SELECTED"],
        [
            ("Shown in the normal work queue\nContinues below for individual review", "Hidden from the normal queue\nRetained until deliberately reselected"),
        ],
        [4680, 4680],
    )
    add_flow_arrow(doc)
    add_heading(doc, "Selected branch - one report at a time", 2)
    add_table(
        doc,
        ["NEW MATTER", "EXISTING MATTER", "DUPLICATE / NO ACTION"],
        [
            (
                "Raise a case\nNo email is sent",
                "Link to the existing case",
                "Resolve the report with a reason",
            ),
        ],
        [3120, 3120, 3120],
    )
    add_flow_arrow(doc)
    add_callout(
        doc,
        "3  CASE CONTROL",
        "Investigate and contact the club  ->  Dave or Warren independently approves another administrator's work  ->  Denver gives final sign-off and issues the outcome.",
        kind="success",
    )
    doc.core_properties.title = "GMCL Ineligible-player Work - Quick Guide"
    doc.core_properties.subject = "Click-by-click staff guide for import selection, live casework, tracker reconciliation and board assurance"
    doc.core_properties.author = "Greater Manchester Cricket League"
    doc.core_properties.keywords = "GMCL, ineligible player, quick guide, Google Form, selection, hidden reports, tracker, board flow"
    doc.core_properties.comments = "compact_reference_guide preset; customer_pack header pattern; August 2026"
    audit_document(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
